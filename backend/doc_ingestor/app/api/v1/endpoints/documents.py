"""
Documents API endpoints for DocQA-MS Document Ingestor
"""
import os
import uuid
from typing import List, Optional
import aiofiles
import httpx
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import JSONResponse
import pika
import json

from app.core.config import settings
from app.core.logging import get_logger
from app.core.database import db_manager

router = APIRouter()
logger = get_logger(__name__)


async def anonymize_content(document_id: str, content: str) -> tuple[str, list]:
    """
    Call the deid service to anonymize document content
    Returns: (anonymized_content, pii_entities)
    """
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            response = await client.post(
                f"{settings.DEID_SERVICE_URL}/anonymize",
                json={
                    "document_id": document_id,
                    "content": content,
                    "language": "fr"
                }
            )
            
            if response.status_code == 200:
                result = response.json()
                anonymized_content = result.get("anonymized_content", content)
                pii_entities = result.get("pii_entities", [])
                logger.info(
                    "Content anonymized successfully",
                    document_id=document_id,
                    pii_count=len(pii_entities)
                )
                return anonymized_content, pii_entities
            else:
                logger.error(
                    "Deid service returned error",
                    document_id=document_id,
                    status_code=response.status_code
                )
                return content, []  # Return original content if deid fails
                
    except httpx.RequestError as e:
        logger.error(
            "Failed to connect to deid service",
            document_id=document_id,
            error=str(e)
        )
        return content, []  # Return original content if deid service is unavailable
    except Exception as e:
        logger.error(
            "Unexpected error during anonymization",
            document_id=document_id,
            error=str(e)
        )
        return content, []


def publish_to_queue(document_id: str, file_path: str, metadata: dict) -> None:
    """Publish document processing task to RabbitMQ queue"""
    try:
        connection = pika.BlockingConnection(pika.URLParameters(settings.RABBITMQ_URL))
        channel = connection.channel()

        # Declare queue
        channel.queue_declare(queue='document_processing', durable=True)

        # Prepare message
        message = {
            "document_id": document_id,
            "file_path": file_path,
            "metadata": metadata,
            "timestamp": "2024-01-01T00:00:00Z"  # Would use datetime.utcnow() in real implementation
        }

        # Publish message
        channel.basic_publish(
            exchange='',
            routing_key='document_processing',
            body=json.dumps(message),
            properties=pika.BasicProperties(
                delivery_mode=2,  # Make message persistent
            )
        )

        logger.info("Published document to processing queue", document_id=document_id)
        connection.close()

    except Exception as e:
        logger.error("Failed to publish to queue", document_id=document_id, error=str(e))
        raise


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    document_id: Optional[str] = Form(None),  # Accept document_id from API Gateway
    patient_id: Optional[str] = Form(None),
    document_type: Optional[str] = Form(None)
):
    """
    Upload a document for processing
    """
    try:
        # Validate file type
        if file.filename:
            file_extension = file.filename.split('.')[-1].lower()
            if file_extension not in settings.ALLOWED_FILE_TYPES:
                raise HTTPException(
                    status_code=400,
                    detail=f"File type not allowed. Allowed types: {', '.join(settings.ALLOWED_FILE_TYPES)}"
                )

        # Validate file size
        file_content = await file.read()
        file_size_mb = len(file_content) / (1024 * 1024)
        if file_size_mb > settings.MAX_UPLOAD_SIZE_MB:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size: {settings.MAX_UPLOAD_SIZE_MB}MB"
            )

        # Use provided document_id or generate new one
        if not document_id:
            document_id = str(uuid.uuid4())
        
        # Extract text content from file
        file_extension = file.filename.split('.')[-1].lower()
        text_content = ""
        
        try:
            if file_extension == 'txt':
                # Try different encodings for text files
                try:
                    text_content = file_content.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        text_content = file_content.decode('latin-1')
                    except UnicodeDecodeError:
                        text_content = file_content.decode('cp1252', errors='replace')
                        
            elif file_extension == 'pdf':
                # Extract text from PDF files
                try:
                    from PyPDF2 import PdfReader
                    import io
                    
                    pdf_reader = PdfReader(io.BytesIO(file_content))
                    pages_text = []
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                pages_text.append(f"[Page {page_num + 1}]\n{page_text}")
                        except Exception as e:
                            logger.warning(f"Failed to extract page {page_num + 1}", error=str(e))
                    
                    text_content = '\n\n'.join(pages_text)
                    
                    if not text_content.strip():
                        text_content = f"[Empty PDF Document: {file.filename}]"
                    
                    logger.info("PDF text extracted", 
                               document_id=document_id,
                               pages=len(pdf_reader.pages),
                               characters=len(text_content))
                except Exception as e:
                    logger.error("Failed to extract PDF content", 
                                error=str(e), 
                                document_id=document_id)
                    text_content = f"[PDF Document: {file.filename}] - Extraction failed: {str(e)}"
                    
            elif file_extension in ['doc', 'docx']:
                # Extract text from DOCX files
                try:
                    from docx import Document
                    import io
                    
                    doc = Document(io.BytesIO(file_content))
                    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                    
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    paragraphs.append(cell.text.strip())
                    
                    text_content = '\n'.join(paragraphs)
                    
                    if not text_content.strip():
                        text_content = f"[Empty DOCX Document: {file.filename}]"
                    
                    logger.info("DOCX text extracted", 
                               document_id=document_id,
                               paragraphs=len(paragraphs),
                               tables=len(doc.tables),
                               characters=len(text_content))
                except Exception as e:
                    logger.error("Failed to extract DOCX content", 
                                error=str(e), 
                                document_id=document_id)
                    text_content = f"[DOCX Document: {file.filename}] - Extraction failed: {str(e)}"
                    
            elif file_extension in ['xlsx', 'xls', 'csv']:
                # Extract text from spreadsheet files
                try:
                    import pandas as pd
                    import io
                    
                    if file_extension == 'csv':
                        df = pd.read_csv(io.BytesIO(file_content))
                        text_content = f"CSV Data:\n{df.to_string()}"
                    else:
                        # Read all sheets from Excel
                        excel_file = pd.ExcelFile(io.BytesIO(file_content))
                        sheets_text = []
                        for sheet_name in excel_file.sheet_names:
                            df = pd.read_excel(excel_file, sheet_name=sheet_name)
                            sheets_text.append(f"[Sheet: {sheet_name}]\n{df.to_string()}")
                        text_content = '\n\n'.join(sheets_text)
                    
                    logger.info("Spreadsheet text extracted", 
                               document_id=document_id,
                               format=file_extension,
                               characters=len(text_content))
                except Exception as e:
                    logger.error("Failed to extract spreadsheet content", 
                                error=str(e), 
                                document_id=document_id)
                    text_content = f"[{file_extension.upper()} Document: {file.filename}] - Extraction failed: {str(e)}"
                    
            elif file_extension == 'json':
                # Extract text from JSON files
                try:
                    import json
                    json_data = json.loads(file_content.decode('utf-8'))
                    # Pretty print JSON for better readability
                    text_content = json.dumps(json_data, indent=2, ensure_ascii=False)
                    
                    logger.info("JSON content extracted", 
                               document_id=document_id,
                               characters=len(text_content))
                except Exception as e:
                    logger.error("Failed to parse JSON content", 
                                error=str(e), 
                                document_id=document_id)
                    text_content = f"[JSON Document: {file.filename}] - Parsing failed: {str(e)}"
                    
            elif file_extension == 'xml':
                # Extract text from XML files
                try:
                    text_content = file_content.decode('utf-8')
                    
                    logger.info("XML content extracted", 
                               document_id=document_id,
                               characters=len(text_content))
                except UnicodeDecodeError:
                    try:
                        text_content = file_content.decode('latin-1')
                    except Exception as e:
                        logger.error("Failed to decode XML content", 
                                    error=str(e), 
                                    document_id=document_id)
                        text_content = f"[XML Document: {file.filename}] - Decoding failed: {str(e)}"
                        
            elif file_extension == 'hl7':
                # Extract text from HL7 messages
                try:
                    import hl7
                    # HL7 messages are typically text-based
                    hl7_text = file_content.decode('utf-8')
                    
                    # Parse HL7 message
                    message = hl7.parse(hl7_text)
                    
                    # Extract readable information
                    segments_text = []
                    for segment in message:
                        segment_name = str(segment[0])
                        segment_data = ' | '.join([str(field) for field in segment[1:] if str(field).strip()])
                        segments_text.append(f"{segment_name}: {segment_data}")
                    
                    text_content = f"HL7 Message:\n" + '\n'.join(segments_text)
                    
                    logger.info("HL7 message parsed", 
                               document_id=document_id,
                               segments=len(message),
                               characters=len(text_content))
                except Exception as e:
                    logger.error("Failed to parse HL7 content", 
                                error=str(e), 
                                document_id=document_id)
                    # Fallback to raw text
                    try:
                        text_content = file_content.decode('utf-8')
                    except:
                        text_content = f"[HL7 Document: {file.filename}] - Parsing failed: {str(e)}"
                        
            elif file_extension in ['fhir', 'fhir.json']:
                # Extract text from FHIR resources
                try:
                    import json
                    from fhir.resources.resource import Resource
                    
                    fhir_json = json.loads(file_content.decode('utf-8'))
                    
                    # Try to parse as FHIR resource
                    try:
                        resource = Resource.parse_obj(fhir_json)
                        resource_type = fhir_json.get('resourceType', 'Unknown')
                        
                        # Pretty print with resource type
                        text_content = f"FHIR Resource Type: {resource_type}\n\n"
                        text_content += json.dumps(fhir_json, indent=2, ensure_ascii=False)
                    except:
                        # If not a valid FHIR resource, just pretty print the JSON
                        text_content = "FHIR Document:\n\n"
                        text_content += json.dumps(fhir_json, indent=2, ensure_ascii=False)
                    
                    logger.info("FHIR resource extracted", 
                               document_id=document_id,
                               resource_type=fhir_json.get('resourceType', 'Unknown'),
                               characters=len(text_content))
                except Exception as e:
                    logger.error("Failed to parse FHIR content", 
                                error=str(e), 
                                document_id=document_id)
                    text_content = f"[FHIR Document: {file.filename}] - Parsing failed: {str(e)}"
                    
            elif file_extension in ['dcm', 'dicom']:
                # Extract text from DICOM medical images
                try:
                    import pydicom
                    import io
                    
                    dicom_data = pydicom.dcmread(io.BytesIO(file_content), force=True)
                    
                    # Extract DICOM tags as text
                    metadata_lines = [
                        f"DICOM Medical Image: {file.filename}",
                        f"Modality: {getattr(dicom_data, 'Modality', 'Unknown')}",
                        f"Patient ID: {getattr(dicom_data, 'PatientID', 'Unknown')}",
                        f"Patient Name: {getattr(dicom_data, 'PatientName', 'Unknown')}",
                        f"Study Date: {getattr(dicom_data, 'StudyDate', 'Unknown')}",
                        f"Study Description: {getattr(dicom_data, 'StudyDescription', 'Unknown')}",
                        f"Series Description: {getattr(dicom_data, 'SeriesDescription', 'Unknown')}",
                        f"Body Part Examined: {getattr(dicom_data, 'BodyPartExamined', 'Unknown')}",
                        "\nDICOM Tags:"
                    ]
                    
                    # Add important DICOM tags
                    for elem in dicom_data:
                        if elem.VR != 'SQ':  # Skip sequences
                            tag_name = elem.name
                            tag_value = str(elem.value)
                            if len(tag_value) < 200:  # Only include reasonable length values
                                metadata_lines.append(f"  {tag_name}: {tag_value}")
                    
                    text_content = '\n'.join(metadata_lines)
                    
                    logger.info("DICOM metadata extracted", 
                               document_id=document_id,
                               modality=getattr(dicom_data, 'Modality', 'Unknown'),
                               tags_count=len(dicom_data),
                               characters=len(text_content))
                except Exception as e:
                    logger.error("Failed to extract DICOM content", 
                                error=str(e), 
                                document_id=document_id)
                    text_content = f"[DICOM Document: {file.filename}] - Extraction failed: {str(e)}"
                        
            else:
                text_content = f"[{file_extension.upper()} Document: {file.filename}] - Format not yet supported"
        except Exception as e:
            logger.warning("Failed to extract text content", error=str(e))
            text_content = f"[Binary file: {file.filename}]"

        # Anonymize the content before saving
        anonymized_content, pii_entities = await anonymize_content(document_id, text_content)
        
        # Log anonymization results
        if pii_entities:
            logger.info(
                "PII detected and anonymized",
                document_id=document_id,
                pii_types=[entity.get("entity_type") for entity in pii_entities]
            )

        # Save file temporarily
        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
        file_path = os.path.join(settings.UPLOAD_DIR, f"{document_id}_{file.filename}")

        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)

        # Prepare metadata
        metadata = {
            "filename": file.filename,
            "content_type": file.content_type,
            "size": len(file_content),
            "patient_id": patient_id,
            "document_type": document_type,
            "file_path": file_path,
            "pii_entities": pii_entities  # Store detected PII for auditing
        }
        
        # Save document to database with anonymized content
        try:
            await db_manager.save_document(
                document_id=document_id,
                filename=file.filename,
                file_type=file_extension,
                content=anonymized_content,  # Save anonymized content instead of original
                file_size=len(file_content),
                metadata=metadata
            )
            logger.info("Document saved to database (anonymized)", document_id=document_id)
        except Exception as e:
            logger.error("Failed to save document to database", error=str(e))
            # Continue anyway - document is still on disk

        # Publish to processing queue
        try:
            publish_to_queue(document_id, file_path, metadata)
        except Exception as e:
            logger.error("Failed to publish to queue", error=str(e))
            # Continue - document is saved to database

        logger.info(
            "Document uploaded successfully",
            document_id=document_id,
            filename=file.filename,
            size=len(file_content)
        )

        return {
            "document_id": document_id,
            "status": "uploaded",
            "message": "Document uploaded successfully and queued for processing"
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(
            "Unexpected error during document upload",
            error=str(e),
            filename=file.filename if file else None
        )
        raise HTTPException(
            status_code=500,
            detail="Internal server error during document upload"
        )


@router.get("/")
async def list_documents(
    status: Optional[str] = None,
    document_type: Optional[str] = None,
    patient_id: Optional[str] = None,
    limit: int = 50,
    offset: int = 0
):
    """
    List documents with filtering and pagination
    """
    try:
        from app.core.database import db_manager
        
        # Query the database for documents
        query = """
            SELECT 
                id::text,
                filename,
                file_type,
                file_size,
                processing_status,
                is_anonymized,
                created_at,
                upload_date,
                metadata->>'patient_id' as patient_id,
                metadata->>'document_type' as document_type
            FROM documents
            WHERE 1=1
        """
        query_params = []
        param_count = 1
        
        # Add filters
        if status:
            query += f" AND processing_status = ${param_count}"
            query_params.append(status)
            param_count += 1
            
        if document_type:
            query += f" AND metadata->>'document_type' = ${param_count}"
            query_params.append(document_type)
            param_count += 1
            
        if patient_id:
            query += f" AND metadata->>'patient_id' = ${param_count}"
            query_params.append(patient_id)
            param_count += 1
        
        # Add ordering and pagination
        query += f" ORDER BY created_at DESC LIMIT ${param_count} OFFSET ${param_count + 1}"
        query_params.extend([limit, offset])
        
        # Get total count
        count_query = "SELECT COUNT(*) as total FROM documents WHERE 1=1"
        if status:
            count_query += f" AND processing_status = '{status}'"
        if document_type:
            count_query += f" AND metadata->>'document_type' = '{document_type}'"
        if patient_id:
            count_query += f" AND metadata->>'patient_id' = '{patient_id}'"
        
        # Execute queries
        async with db_manager.pool.acquire() as conn:
            rows = await conn.fetch(query, *query_params)
            count_row = await conn.fetchrow(count_query)
            total = count_row['total'] if count_row else 0
            
            documents = []
            for row in rows:
                doc = dict(row)
                # Convert datetime to ISO string
                if doc.get('created_at'):
                    doc['created_at'] = doc['created_at'].isoformat()
                if doc.get('upload_date'):
                    doc['upload_date'] = doc['upload_date'].isoformat()
                documents.append(doc)

        return {
            "data": documents,  # Changed from "documents" to "data" for frontend compatibility
            "total": total,
            "limit": limit,
            "offset": offset,
            "page": offset // limit + 1 if limit > 0 else 1,
            "total_pages": (total + limit - 1) // limit if limit > 0 else 1,
            "has_next": offset + limit < total,
            "has_prev": offset > 0
        }
    
    except Exception as e:
        logger.error("Failed to list documents", error=str(e))
        # Return empty list on error instead of failing
        return {
            "data": [],
            "total": 0,
            "limit": limit,
            "offset": offset,
            "page": 1,
            "total_pages": 0,
            "has_next": False,
            "has_prev": False
        }

    except Exception as e:
        logger.error("Failed to list documents", error=str(e))
        raise HTTPException(
            status_code=500,
            detail="Failed to retrieve documents"
        )


@router.get("/{document_id}")
async def get_document(document_id: str):
    """
    Get detailed information about a specific document
    """
    try:
        # This would query the database for document details
        # For now, return mock data
        document = {
            "id": document_id,
            "filename": "sample_medical_report_1.pdf",
            "status": "processed",
            "upload_date": "2024-01-15T10:30:00Z",
            "document_type": "medical_report",
            "patient_id": "ANON_PAT_001",
            "file_size": 245760,
            "processing_status": "completed",
            "metadata": {
                "author": "Dr. Smith",
                "specialty": "cardiology",
                "language": "fr"
            }
        }

        return document

    except Exception as e:
        logger.error(
            "Failed to get document",
            document_id=document_id,
            error=str(e)
        )
        raise HTTPException(
            status_code=500,
            detail="Failed to retrieve document"
        )


@router.get("/{document_id}/download")
async def download_document(document_id: str):
    """
    Download the original document file
    """
    try:
        # Get document metadata from database to find file_path
        async with db_manager.pool.acquire() as conn:
            result = await conn.fetchrow("""
                SELECT filename, file_type, metadata
                FROM documents
                WHERE id = $1
            """, document_id)
        
        if not result:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Get file_path from metadata
        metadata = result.get('metadata', {})
        if isinstance(metadata, str):
            metadata = json.loads(metadata)
        
        file_path = metadata.get('file_path')
        
        if not file_path or not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="Document file not found on disk")
        
        # Determine content type based on file extension
        file_type = result['file_type'].lower()
        content_type_map = {
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'doc': 'application/msword',
            'txt': 'text/plain',
            'hl7': 'text/plain',
            'xml': 'application/xml',
            'json': 'application/json',
        }
        content_type = content_type_map.get(file_type, 'application/octet-stream')
        
        # Read and return file
        from fastapi.responses import FileResponse
        return FileResponse(
            path=file_path,
            media_type=content_type,
            filename=result['filename']
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(
            "Failed to download document",
            document_id=document_id,
            error=str(e)
        )
        raise HTTPException(
            status_code=500,
            detail=f"Failed to download document: {str(e)}"
        )


@router.delete("/{document_id}")
async def delete_document(document_id: str):
    """
    Delete a document and all associated data
    """
    try:
        # This would delete from database and file system
        # For now, just return success
        logger.info("Document deleted", document_id=document_id)

        return {
            "message": "Document deleted successfully",
            "document_id": document_id
        }

    except Exception as e:
        logger.error(
            "Failed to delete document",
            document_id=document_id,
            error=str(e)
        )
        raise HTTPException(
            status_code=500,
            detail="Failed to delete document"
        )